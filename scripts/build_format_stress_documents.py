"""Genera documentos ficticios multiformato para pruebas locales de Omega.

No contiene operaciones, personas o empresas reales. Los PDF escaneados están hechos
como imagen intencionalmente, para ejercitar OCR sin depender de fuentes externas.
"""

from __future__ import annotations

import csv
import textwrap
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont


ROOT = Path.cwd() / "corpus-prueba-formatos-extremos"

PEOPLE = [
    "Andrea Córdova Reyes", "Bruno Alvarado Cruz", "Carla Mena Soto", "Diego Rivas León",
    "Estela Pineda Mora", "Fernando Quiroz Gil", "Gloria Nájera Solís", "Hugo Treviño Lara",
]
CITIES = ["Ciudad de México", "Guadalajara", "Monterrey", "Puebla", "Querétaro", "Mérida"]


def ensure(folder: str) -> Path:
    path = ROOT / folder
    path.mkdir(parents=True, exist_ok=True)
    return path


def narrative(index: int, subject: str) -> list[str]:
    person = PEOPLE[index % len(PEOPLE)]
    city = CITIES[index % len(CITIES)]
    return [
        f"El expediente {subject} se abrió para registrar una operación ficticia en {city}. "
        f"La persona responsable, {person}, verificó que los datos de identificación, fecha, "
        "importe y referencia documental fueran consistentes antes de autorizar el siguiente paso.",
        "La información se conserva con acceso restringido. Cualquier solicitud de un tercero "
        "debe validar identidad, facultades y propósito antes de compartir documentos, montos o "
        "datos personales. Si falta evidencia, el equipo debe registrar una observación y pedir "
        "la documentación necesaria, sin inventar una conclusión.",
        "El seguimiento incluye confirmar requisitos, conservar la evidencia de la decisión y "
        "revisar las fechas límite. Una diferencia entre el registro y el documento de origen "
        "debe escalarse al responsable de cumplimiento antes de cerrar el expediente.",
    ]


def fields(index: int, type_name: str, amount: int) -> list[tuple[str, str]]:
    return [
        ("Folio", f"FMT-26-{index + 1:04d}"),
        ("Tipo de documento", type_name),
        ("Estado", ["Activo", "En revisión", "Cerrado"][index % 3]),
        ("Ciudad base", CITIES[index % len(CITIES)]),
        ("Responsable", PEOPLE[(index + 2) % len(PEOPLE)]),
        ("Fecha de registro", f"2026-0{1 + (index % 8)}-{5 + index:02d}"),
        ("Importe total", f"${amount:,.2f} MXN"),
    ]


def report_fonts() -> tuple[str, str]:
    """Embebe una fuente TrueType para que la renderización PDF sea consistente."""
    regular = next(
        (Path(candidate) for candidate in [
            "/System/Library/Fonts/Supplemental/Arial.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        ] if Path(candidate).is_file()),
        None,
    )
    bold = next(
        (Path(candidate) for candidate in [
            "/System/Library/Fonts/Supplemental/Arial Bold.ttf",
            "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
        ] if Path(candidate).is_file()),
        regular,
    )
    if regular is None:
        return "Helvetica", "Helvetica-Bold"
    pdfmetrics.registerFont(TTFont("OmegaStressRegular", str(regular)))
    pdfmetrics.registerFont(TTFont("OmegaStressBold", str(bold)))
    return "OmegaStressRegular", "OmegaStressBold"


def pdf_text(index: int) -> None:
    folder = ensure("01_pdf_texto")
    type_name = ["Contrato de suministro", "Informe de auditoría", "Orden de servicio"][index % 3]
    amount = 18000 + index * 2400
    path = folder / f"{index + 1:03d}-{type_name.lower().replace(' ', '-')}.pdf"
    regular_font, bold_font = report_fonts()
    styles = getSampleStyleSheet()
    body = ParagraphStyle("OmegaBody", parent=styles["BodyText"], fontName=regular_font, fontSize=10.5, leading=14, spaceAfter=9)
    title = ParagraphStyle("OmegaTitle", parent=styles["Title"], fontName=bold_font, fontSize=18, leading=22, textColor=colors.HexColor("#1F4D78"))
    table_label = ParagraphStyle("OmegaTableLabel", parent=body, fontName=bold_font, fontSize=9.5, leading=12, spaceAfter=0)
    table_value = ParagraphStyle("OmegaTableValue", parent=body, fontName=regular_font, fontSize=9.5, leading=12, spaceAfter=0)
    story = [Paragraph(type_name, title), Spacer(1, 0.15 * inch)]
    data = [[Paragraph(label, table_label), Paragraph(value, table_value)] for label, value in fields(index, type_name, amount)]
    table = Table(data, colWidths=[1.65 * inch, 4.65 * inch])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#E8EEF5")),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#C8D0D9")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 7),
        ("RIGHTPADDING", (0, 0), (-1, -1), 7),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    story += [table, Spacer(1, 0.22 * inch), Paragraph("Antecedentes y controles", styles["Heading2"])]
    for paragraph in narrative(index, type_name.lower()):
        story.append(Paragraph(paragraph, body))
    story.append(Paragraph("Documento sintético creado exclusivamente para pruebas de formato y recuperación local.", body))
    SimpleDocTemplate(str(path), pagesize=letter, leftMargin=0.8 * inch, rightMargin=0.8 * inch, topMargin=0.72 * inch, bottomMargin=0.72 * inch).build(story)


def get_scan_font(size: int):
    for candidate in [
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Verdana.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]:
        if Path(candidate).is_file():
            return ImageFont.truetype(candidate, size)
    return ImageFont.load_default()


def pdf_scan(index: int) -> None:
    folder = ensure("02_pdf_escaneado_ocr")
    type_name = ["Recibo de inspección", "Acta de entrega"][index % 2]
    amount = 7400 + index * 1100
    path = folder / f"{index + 1:03d}-escaneo-{type_name.lower().replace(' ', '-')}.pdf"
    image = Image.new("RGB", (1700, 2200), "#FBF8F0")
    draw = ImageDraw.Draw(image)
    title_font = get_scan_font(46)
    label_font = get_scan_font(27)
    body_font = get_scan_font(25)
    draw.text((120, 100), type_name.upper(), fill="#1B365D", font=title_font)
    y = 220
    for label, value in fields(index + 30, type_name, amount):
        draw.text((125, y), f"{label}: {value}", fill="#171717", font=label_font)
        y += 54
    y += 48
    draw.line((120, y, 1580, y), fill="#8B8B8B", width=2)
    y += 52
    for heading, paragraph in zip(["OBSERVACIONES", "PRIVACIDAD Y SEGUIMIENTO", "CIERRE"], narrative(index + 30, type_name.lower())):
        draw.text((125, y), heading, fill="#1B365D", font=label_font)
        y += 42
        for line in textwrap.wrap(paragraph, width=88):
            draw.text((125, y), line, fill="#202020", font=body_font)
            y += 34
        y += 28
    image.save(path, "PDF", resolution=160.0)


def word_doc(index: int) -> None:
    folder = ensure("03_word_docx")
    type_name = ["Minuta de comité", "Manual operativo", "Solicitud de compra", "Reporte de incidente"][index % 4]
    amount = 21500 + index * 1800
    path = folder / f"{index + 1:03d}-{type_name.lower().replace(' ', '-')}.docx"
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run(type_name)
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(31, 77, 120)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in fields(index + 50, type_name, amount):
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        for run in cells[0].paragraphs[0].runs:
            run.bold = True
    for heading, paragraph in zip(["Antecedentes", "Controles de información", "Seguimiento"], narrative(index + 50, type_name.lower())):
        doc.add_heading(heading, level=1)
        doc.add_paragraph(paragraph)
    doc.add_paragraph("Archivo ficticio para validar extracción de Word, campos tabulares, encabezados y texto largo.")
    doc.save(path)


def markdown_docs() -> None:
    folder = ensure("06_markdown_largo")
    for index in range(10):
        type_name = ["Política de conservación", "Procedimiento de aprobación"][index % 2]
        amount = 9900 + index * 990
        field_lines = "\n".join(f"{label}: {value}" for label, value in fields(index + 70, type_name, amount))
        content = f"# {type_name}\n\n{field_lines}\n\n## Alcance\n\n" + "\n\n".join(narrative(index + 70, type_name.lower())) + "\n\n## Nota de prueba\n\nEste Markdown extenso es ficticio y prueba rutas, filtros, números, fechas y evidencia textual.\n"
        (folder / f"{index + 1:03d}-{type_name.lower().replace(' ', '-')}.md").write_text(content, encoding="utf-8")


def csv_docs() -> None:
    folder = ensure("05_csv")
    for index in range(6):
        path = folder / f"{index + 1:03d}-movimientos-operativos.csv"
        with path.open("w", newline="", encoding="utf-8") as stream:
            writer = csv.writer(stream)
            writer.writerow(["Folio", "Tipo de documento", "Estado", "Ciudad base", "Fecha de registro", "Importe total", "Observación"])
            for row in range(8):
                sequence = 100 + index * 8 + row
                writer.writerow([
                    f"CSV-26-{sequence:04d}", "Movimiento operativo", ["Activo", "Cerrado"][row % 2],
                    CITIES[(index + row) % len(CITIES)], f"2026-0{1 + (row % 8)}-{10 + row:02d}",
                    1250 + sequence * 11,
                    "Registro sintético; validar origen, importe y estado antes de cerrar.",
                ])


def problematic_files() -> None:
    folder = ensure("07_archivos_problematicos")
    (folder / "001-archivo-vacio.txt").write_bytes(b"")
    (folder / "002-pdf-truncado.pdf").write_text("Este archivo simula un PDF incompleto y no debe bloquear la indexación.", encoding="utf-8")
    (folder / "003-extension-engañosa.docx").write_text("No soy un DOCX válido; Omega debe informar el error sin caer.", encoding="utf-8")
    content = "# Archivo con caracteres y nombre largo\n\nFolio: EDGE-26-0001\nEstado: En revisión\n\nContenido ficticio para comprobar rutas, Unicode y nombres extensos.\n"
    (folder / "004-ñandú-área-operativa-con-nombre-muy-largo-y-espacios.md").write_text(content, encoding="utf-8")


def main() -> None:
    for folder in [
        "01_pdf_texto", "02_pdf_escaneado_ocr", "03_word_docx", "04_excel_xlsx",
        "05_csv", "06_markdown_largo", "07_archivos_problematicos",
    ]:
        ensure(folder)
    for index in range(12):
        pdf_text(index)
    for index in range(8):
        pdf_scan(index)
    for index in range(8):
        word_doc(index)
    markdown_docs()
    csv_docs()
    problematic_files()
    print(f"Base creada en {ROOT}")


if __name__ == "__main__":
    main()
